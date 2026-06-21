"""Convert the project Markdown reports to clean, professional Word (.docx).
Tailored to the Markdown constructs used in these reports (ATX headings,
bold, bullet/numbered lists, GitHub pipe tables, rules, blockquotes)."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent.parent
JOBS = [
    (BASE / "06_Reports/A_Methods_Decision_Log.md",
     BASE / "06_Reports/A_Methods_Decision_Log.docx"),
    (BASE / "06_Reports/B_Final_Method_Stepwise.md",
     BASE / "06_Reports/B_Final_Method_Stepwise.docx"),
    (BASE / "04_Submission/Stage2_Progress_AquaCascade.md",
     BASE / "04_Submission/Stage2_Progress_AquaCascade.docx"),
    (BASE / "04_Submission/ExperimentLogs/"
     "ExperimentLog_01_Cascade_Hypothesis_Rejected.md",
     BASE / "04_Submission/ExperimentLogs/"
     "ExperimentLog_01_Cascade_Hypothesis_Rejected.docx"),
    (BASE / "04_Submission/ExperimentLogs/"
     "ExperimentLog_02_Signature_Methodology_By_Evidence.md",
     BASE / "04_Submission/ExperimentLogs/"
     "ExperimentLog_02_Signature_Methodology_By_Evidence.docx"),
    (BASE / "04_Submission/ExperimentLogs/"
     "ExperimentLog_03_Production_Triage_State_Generalization_Limit.md",
     BASE / "04_Submission/ExperimentLogs/"
     "ExperimentLog_03_Production_Triage_State_Generalization_Limit.docx"),
]
NAVY = RGBColor(0x1B, 0x3A, 0x5B)
HEADERFILL = "1B3A5B"
CHARTS = BASE / "03_Outputs" / "Charts"

# (anchor substring in a heading) -> [(png, caption), ...]
IMAGES = {
    "A_Methods_Decision_Log.md": {
        "Every Method Considered": [("experiment_map.png",
            "Every method tested, decided by evidence (overview).")],
        "## 1.": [("spatial_cascade_results.png",
            "County regional-risk term: base 0.695 -> 0.742 (+0.047).")],
        "## 2.": [("polygon_cascade_results.png",
            "True polygon-adjacency cascade: rejected (sparse graph).")],
        "## 3.": [("unknown_triage_results.png",
            "First triage model and investigation-targeting.")],
        "## 4.": [("model_signature_results.png",
            "Signature risk model (leakage-corrected, ROC 0.683).")],
        "## 5.": [("sig_depth_ab.png",
            "Signature depth 2 vs 3 vs log (25 paired CV splits).")],
        "## 6.": [("sig_calculus_ab.png",
            "Ito vs Stratonovich (25 paired CV splits).")],
        "## 7.": [("process_diagnostics.png",
            "Process diagnostics: drift-dominated, not Brownian.")],
        "## 8.": [("lsl_optimizer_results.png",
            "Replacement optimization (real edge 4.2%).")],
        "## 9.": [("triage_production.png",
            "Selected method: calibration, PR, and lift.")],
        "## 10.": [("leadpath_signature_ab.png",
            "Lead-sample PB90 path signature on triage: rejected.")],
    },
    "B_Final_Method_Stepwise.md": {
        "Complete Step-by-Step": [("experiment_map.png",
            "Project overview: every method tested and its decision.")],
        "PART II —": [("model_signature_results.png",
            "Track-1 signature model: ROC and top LASSO predictors."),
            ("spatial_cascade_results.png",
            "County regional-risk term effect (+0.047 ROC).")],
        "PART III —": [("sig_depth_ab.png",
            "Signature depth A/B (25 paired CV splits)."),
            ("sig_calculus_ab.png",
            "Ito vs Stratonovich A/B."),
            ("process_diagnostics.png",
            "Stochastic-nature diagnostics.")],
        "PART IV —": [("method_flow.png",
            "The selected method, step by step."),
            ("triage_production.png",
            "Calibration, precision-recall, and investigation lift.")],
        "PART V —": [("leadpath_signature_ab.png",
            "Lead-sample PB90 path signature on the triage target.")],
        "PART VI —": [("polygon_cascade_results.png",
            "Inter-utility polygon cascade (rejected)."),
            ("lsl_optimizer_results.png",
            "Replacement cost-optimization (de-prioritized).")],
    },
    "Stage2_Progress_AquaCascade.md": {
        "What I built": [("method_flow.png",
            "The pipeline, step by step.")],
        "Experiment log": [("experiment_map.png",
            "Every hypothesis tested and the evidence-based decision.")],
        "Headline verified result": [("triage_production.png",
            "Calibrated triage model: calibration, PR, and lift.")],
    },
}


def base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Arial"
    st.font.size = Pt(11)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    for name, sz in (("Heading 1", 18), ("Heading 2", 14),
                     ("Heading 3", 12)):
        s = doc.styles[name]
        s.font.name = "Arial"
        s.font.size = Pt(sz)
        s.font.bold = True
        s.font.color.rgb = NAVY
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(1))


def shade(cell, hexfill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


def add_inline(par, text):
    # split on ** (bold) and `code`
    for chunk in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            r = par.add_run(chunk[2:-2])
            r.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = par.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            par.add_run(chunk)


def emit_table(doc, rows):
    hdr = [c.strip() for c in rows[0].strip().strip("|").split("|")]
    body = [[c.strip() for c in r.strip().strip("|").split("|")]
            for r in rows[2:]]
    t = doc.add_table(rows=1, cols=len(hdr))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    for i, h in enumerate(hdr):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        add_inline(p, h)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, HEADERFILL)
    for brow in body:
        cells = t.add_row().cells
        for i, val in enumerate(brow[:len(hdr)]):
            cells[i].text = ""
            add_inline(cells[i].paragraphs[0], val)
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    if r.font.size is None:
                        r.font.size = Pt(9.5)


def add_figure(doc, png, caption):
    path = CHARTS / png
    if not path.exists():
        return
    import struct
    with open(path, "rb") as fh:
        fh.read(16)
        w, h = struct.unpack(">II", fh.read(8))
    aspect = w / h
    max_w, max_h = 6.4, 3.5
    wi = max_w
    hi = wi / aspect
    if hi > max_h:
        hi = max_h
        wi = hi * aspect
    doc.add_paragraph()
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(path), width=Inches(wi))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("Figure. " + caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x5B, 0x6B, 0x7B)


def inject(doc, fname, raw_line):
    for anchor, figs in IMAGES.get(fname, {}).items():
        if anchor in raw_line:
            for png, cap in figs:
                add_figure(doc, png, cap)
            break


def convert(md_path, out_path):
    fname = md_path.name
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    base_style(doc)
    i = 0
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()
        if not s:
            i += 1
            continue
        if s.startswith("|") and i + 1 < len(lines) and \
                re.match(r"^\|[\s:|-]+\|?$", lines[i + 1].strip()):
            blk = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                blk.append(lines[i])
                i += 1
            emit_table(doc, blk)
            doc.add_paragraph()
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
            inject(doc, fname, s)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
            inject(doc, fname, s)
        elif s.startswith("# "):
            h = doc.add_heading(s[2:], level=1)
            h.runs[0].font.size = Pt(20)
            inject(doc, fname, s)
        elif s in ("---", "***", "___"):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            bd = OxmlElement("w:pBdr")
            b = OxmlElement("w:bottom")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "6")
            b.set(qn("w:color"), HEADERFILL)
            bd.append(b)
            pPr.append(bd)
        elif s.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(s[2:])
            r.italic = True
        elif re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s", "", s))
        elif s.startswith("- "):
            indent = len(ln) - len(ln.lstrip())
            p = doc.add_paragraph(style="List Bullet"
                                  if indent < 2 else "List Bullet 2")
            add_inline(p, s[2:])
        else:
            p = doc.add_paragraph()
            add_inline(p, s)
        i += 1
    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    for src, dst in JOBS:
        convert(src, dst)
        print("wrote", dst.name, dst.stat().st_size, "bytes")
